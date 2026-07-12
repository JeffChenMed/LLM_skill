from __future__ import annotations

import argparse
import copy
import re
import shutil
import subprocess
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Emu, Pt
from PIL import Image

REF_RE = re.compile(
    r"\b(?:Supplementary\s+)?(?:Table|Tables|Figure|Figures|Fig\.|Figs\.)\s+"
    r"\d+(?:[A-Za-z])?(?:\s*(?:,|and|-)\s*\d+(?:[A-Za-z])?)*"
)
LEGEND_START_RE = re.compile(r"^(?:Fig\.|Figure)\s+\d+\.")


def set_run_font(run, name: str = "Times New Roman", size_pt: float = 12.0) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def remove_paragraph(paragraph) -> None:
    elem = paragraph._element
    elem.getparent().remove(elem)
    paragraph._p = paragraph._element = None


def append_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def stop_before_references(doc: Document) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().upper() == "REFERENCES":
            return i
    return len(doc.paragraphs)


def force_introduction_page_break(doc: Document) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().upper() == "INTRODUCTION":
            paragraph.paragraph_format.page_break_before = True
            return True
    return False


def effective_line_spacing(paragraph) -> float | None:
    value = paragraph.paragraph_format.line_spacing
    style = paragraph.style
    while value is None and style is not None:
        value = style.paragraph_format.line_spacing
        style = style.base_style
    return float(value) if isinstance(value, (int, float)) else value


def copy_run_format(src, dst) -> None:
    if src.style:
        dst.style = src.style
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    dst.font.name = src.font.name
    dst.font.size = src.font.size
    dst.font.bold = src.font.bold
    dst.font.italic = src.font.italic
    dst.font.underline = src.font.underline
    dst.font.superscript = src.font.superscript
    dst.font.subscript = src.font.subscript
    dst.font.small_caps = src.font.small_caps
    dst.font.all_caps = src.font.all_caps
    dst.font.strike = src.font.strike
    dst.font.double_strike = src.font.double_strike
    dst.font.color.rgb = src.font.color.rgb
    dst.font.highlight_color = src.font.highlight_color
    if src._element.rPr is not None:
        text = dst.text
        dst._element.clear_content()
        dst._element.insert(0, copy.deepcopy(src._element.rPr))
        dst.text = text


def bold_refs_in_paragraph(paragraph) -> list[str]:
    text = paragraph.text
    matches = list(REF_RE.finditer(text))
    if not matches:
        return []

    force = [False] * len(text)
    for match in matches:
        for idx in range(match.start(), match.end()):
            force[idx] = True

    char_sources = []
    for run in paragraph.runs:
        for ch in run.text:
            char_sources.append((ch, run))
    if len(char_sources) != len(text):
        raise RuntimeError("Run/text length mismatch while bolding references.")

    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    i = 0
    while i < len(text):
        src = char_sources[i][1]
        is_forced = force[i]
        j = i + 1
        while j < len(text) and char_sources[j][1] is src and force[j] == is_forced:
            j += 1
        new_run = paragraph.add_run(text[i:j])
        copy_run_format(src, new_run)
        if is_forced:
            new_run.bold = True
            new_run.font.bold = True
        i = j

    return [m.group(0) for m in matches]


def bold_in_text_citations(doc: Document) -> list[tuple[int, list[str]]]:
    changed = []
    stop = stop_before_references(doc)
    for i, paragraph in enumerate(doc.paragraphs[:stop]):
        refs = bold_refs_in_paragraph(paragraph)
        if refs:
            changed.append((i, refs))
    return changed


def extract_centralized_legends(doc: Document) -> list[tuple[str, str]]:
    start_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() == "figure legends":
            start_idx = i
            break
    if start_idx is None:
        return []

    blocks: list[list[str]] = []
    current: list[str] | None = None
    for paragraph in doc.paragraphs[start_idx + 1 :]:
        text = paragraph.text.strip()
        if not text:
            continue
        if LEGEND_START_RE.match(text):
            if current:
                blocks.append(current)
            current = [text]
        elif current is not None:
            current.append(text)
    if current:
        blocks.append(current)

    legends = []
    for block in blocks:
        joined = "\n".join(block).strip()
        if "\n" in joined:
            title, body = joined.split("\n", 1)
        else:
            title, body = joined, ""
        legends.append((title.strip(), body.strip()))
    return legends


def export_pdf_to_jpg(pdf_path: Path, out_dir: Path, poppler_pdftoppm: Path | None, dpi: int) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_prefix = out_dir / pdf_path.stem
    out_path = out_prefix.with_suffix(".jpg")
    if out_path.exists():
        out_path.unlink()
    if not poppler_pdftoppm:
        raise RuntimeError("PDF figure export requires --pdftoppm pointing to Poppler pdftoppm.")
    subprocess.run(
        [str(poppler_pdftoppm), "-jpeg", "-r", str(dpi), "-f", "1", "-singlefile", str(pdf_path), str(out_prefix)],
        check=True,
    )
    if not out_path.exists():
        raise RuntimeError(f"Could not export {pdf_path} to JPG.")
    return out_path


def fit_image_width(image_path: Path, max_width_emu: int, max_height_emu: int) -> Emu:
    with Image.open(image_path) as im:
        px_w, px_h = im.size
    ratio = min(max_width_emu / px_w, max_height_emu / px_h)
    return Emu(int(px_w * ratio))


def rebuild_figure_blocks(
    doc: Document,
    figure_images: list[Path],
    legends: list[tuple[str, str]],
    body_pt: float,
    font_name: str,
    line_spacing: float,
    image_height_fraction: float,
) -> None:
    if len(figure_images) != len(legends):
        raise RuntimeError(f"Figure count ({len(figure_images)}) and legend count ({len(legends)}) differ.")

    last_ref_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == "EndNote Bibliography" and paragraph.text.strip():
            last_ref_idx = i
    if last_ref_idx is None:
        raise RuntimeError("Could not find reference bibliography paragraphs.")

    for paragraph in list(doc.paragraphs)[last_ref_idx + 1 :]:
        remove_paragraph(paragraph)

    section = doc.sections[-1]
    max_width = int(section.page_width - section.left_margin - section.right_margin)
    usable_height = int(section.page_height - section.top_margin - section.bottom_margin)
    max_image_height = int(usable_height * image_height_fraction)

    append_page_break(doc)
    for idx, (image_path, (title, body)) in enumerate(zip(figure_images, legends)):
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.space_before = Pt(0)
        img_para.paragraph_format.space_after = Pt(6)
        img_para.add_run().add_picture(str(image_path), width=fit_image_width(image_path, max_width, max_image_height))

        title_para = doc.add_paragraph()
        title_para.paragraph_format.line_spacing = line_spacing
        title_para.paragraph_format.space_before = None
        title_para.paragraph_format.space_after = None
        title_run = title_para.add_run(title)
        title_run.bold = True
        set_run_font(title_run, font_name, body_pt)

        if body:
            body_para = doc.add_paragraph()
            body_para.paragraph_format.line_spacing = line_spacing
            body_para.paragraph_format.space_before = None
            body_para.paragraph_format.space_after = None
            body_run = body_para.add_run(body)
            body_run.bold = False
            set_run_font(body_run, font_name, body_pt)

        if idx != len(figure_images) - 1:
            append_page_break(doc)


def audit_doc(
    doc: Document,
    *,
    require_introduction_break: bool = False,
    expected_legend_line_spacing: float | None = None,
) -> list[str]:
    issues = []
    if require_introduction_break and not any(
        p.text.strip().upper() == "INTRODUCTION" and p.paragraph_format.page_break_before
        for p in doc.paragraphs
    ):
        issues.append("INTRODUCTION does not have explicit page_break_before.")

    stop = stop_before_references(doc)
    for i, paragraph in enumerate(doc.paragraphs[:stop]):
        text = paragraph.text
        for match in REF_RE.finditer(text):
            pos = 0
            bolds = []
            for run in paragraph.runs:
                start, end = pos, pos + len(run.text)
                pos = end
                if end > match.start() and start < match.end():
                    bolds.append(run.bold is True or run.font.bold is True)
            if not bolds or not all(bolds):
                issues.append(f"Unbold citation at paragraph {i}: {match.group(0)}")

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if LEGEND_START_RE.match(text):
            if not all((r.bold is True or r.font.bold is True) for r in paragraph.runs if r.text.strip()):
                issues.append(f"Legend title not fully bold at paragraph {i}: {text[:50]}")
            if (
                expected_legend_line_spacing is not None
                and effective_line_spacing(paragraph) != expected_legend_line_spacing
            ):
                issues.append(
                    f"Legend title line spacing does not match {expected_legend_line_spacing} "
                    f"at paragraph {i}."
                )
    return issues


def main() -> None:
    parser = argparse.ArgumentParser(description="Apply CZF manuscript submission formatting fixes.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path, help="Output DOCX. Defaults to in-place.")
    parser.add_argument("--figure-pdf", action="append", type=Path, default=[], help="Figure PDF to export and insert, in order.")
    parser.add_argument("--figure-jpg", action="append", type=Path, default=[], help="Figure JPG to insert, in order.")
    parser.add_argument("--pdftoppm", type=Path, help="Path to Poppler pdftoppm executable for PDF figure export.")
    parser.add_argument("--jpg-dir", type=Path, default=Path("figure_jpg_czf"))
    parser.add_argument("--body-pt", type=float, default=12.0)
    parser.add_argument("--font", default="Times New Roman")
    parser.add_argument(
        "--line-spacing",
        type=float,
        help="Legend line-spacing multiple. Defaults to the target document Normal style.",
    )
    parser.add_argument("--image-height-fraction", type=float, default=0.93)
    parser.add_argument("--force-introduction-page-break", action="store_true")
    parser.add_argument("--audit-only", action="store_true")
    args = parser.parse_args()

    out = args.out or args.docx
    if out == args.docx and not args.audit_only:
        backup = args.docx.with_suffix(".before_czf_format.docx")
        if not backup.exists():
            shutil.copy2(args.docx, backup)

    doc = Document(args.docx)
    line_spacing = args.line_spacing
    if line_spacing is None:
        value = doc.styles["Normal"].paragraph_format.line_spacing
        line_spacing = float(value) if isinstance(value, (int, float)) else 1.5
    if not args.audit_only:
        if args.force_introduction_page_break:
            force_introduction_page_break(doc)
        bold_in_text_citations(doc)

        figure_images = list(args.figure_jpg)
        if args.figure_pdf:
            for pdf_path in args.figure_pdf:
                figure_images.append(export_pdf_to_jpg(pdf_path, args.jpg_dir, args.pdftoppm, dpi=300))

        if figure_images:
            legends = extract_centralized_legends(doc)
            rebuild_figure_blocks(
                doc,
                figure_images,
                legends,
                args.body_pt,
                args.font,
                line_spacing,
                args.image_height_fraction,
            )
        doc.save(out)
        doc = Document(out)

    issues = audit_doc(
        doc,
        require_introduction_break=args.force_introduction_page_break,
        expected_legend_line_spacing=line_spacing,
    )
    if issues:
        print("AUDIT ISSUES")
        for issue in issues:
            print(f"- {issue}")
        raise SystemExit(1)
    print("AUDIT OK")


if __name__ == "__main__":
    main()
