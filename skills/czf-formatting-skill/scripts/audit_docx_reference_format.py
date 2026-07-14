#!/usr/bin/env python3
"""Compare a target DOCX with a controlling reference DOCX format contract."""

from __future__ import annotations

import argparse
import json
import re
import sys
import xml.etree.ElementTree as ET
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.oxml.ns import qn


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--reference", type=Path, required=True)
    parser.add_argument("--target", type=Path, required=True)
    parser.add_argument("--expected-figures", type=int)
    parser.add_argument(
        "--require-section-order",
        help="Exact paragraph labels separated by '>', in required document order.",
    )
    parser.add_argument(
        "--require-paragraph-prefix-order",
        help="Paragraph prefixes separated by '>', in required document order.",
    )
    parser.add_argument(
        "--require-page-break-before",
        action="append",
        default=[],
        help="Exact paragraph label that must have page-break-before; repeat as needed.",
    )
    parser.add_argument(
        "--figures-after-heading",
        help="Require every visible body drawing to occur after this exact heading.",
    )
    parser.add_argument(
        "--require-figure-legend-pairs",
        action="store_true",
        help="Require a visible Figure/Fig. legend after each drawing and before the next drawing.",
    )
    parser.add_argument("--json", type=Path)
    return parser.parse_args()


def mode(values: Iterable[Any]) -> Any:
    values = [value for value in values if value is not None]
    return Counter(values).most_common(1)[0][0] if values else None


def pt(value) -> float | None:
    return None if value is None else round(value.pt, 3)


def inches(value) -> float | None:
    return None if value is None else round(value.inches, 4)


def style_value(style, attribute: str, *, font: bool = False):
    while style is not None:
        container = style.font if font else style.paragraph_format
        value = getattr(container, attribute)
        if value is not None:
            return value
        style = style.base_style
    return None


def paragraph_value(paragraph, attribute: str):
    value = getattr(paragraph.paragraph_format, attribute)
    return value if value is not None else style_value(paragraph.style, attribute)


def run_font(run, paragraph, attribute: str):
    value = getattr(run.font, attribute)
    if value is not None:
        return value
    if run.style is not None:
        value = style_value(run.style, attribute, font=True)
        if value is not None:
            return value
    return style_value(paragraph.style, attribute, font=True)


def weighted_run_mode(paragraphs, attribute: str):
    weighted = []
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            value = run_font(run, paragraph, attribute)
            if attribute == "size":
                value = pt(value)
            weighted.extend([value] * max(1, len(run.text)))
    return mode(weighted)


def is_heading(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or len(text) > 90:
        return False
    if paragraph.style.name.startswith("Heading"):
        return True
    letters = [character for character in text if character.isalpha()]
    return bool(letters) and text.upper() == text and any(
        run.text.strip() and (run.bold is True or run.font.bold is True)
        for run in paragraph.runs
    )


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def section_signature(doc: Document) -> dict[str, Any]:
    section = doc.sections[0]
    sect_pr = section._sectPr
    return {
        "page_width_in": inches(section.page_width),
        "page_height_in": inches(section.page_height),
        "top_margin_in": inches(section.top_margin),
        "right_margin_in": inches(section.right_margin),
        "bottom_margin_in": inches(section.bottom_margin),
        "left_margin_in": inches(section.left_margin),
        "header_distance_in": inches(section.header_distance),
        "footer_distance_in": inches(section.footer_distance),
        "line_numbering": sect_pr.find(qn("w:lnNumType")) is not None,
    }


def paragraph_group_signature(paragraphs) -> dict[str, Any] | None:
    paragraphs = [paragraph for paragraph in paragraphs if paragraph.text.strip()]
    if not paragraphs:
        return None
    line_spacing = []
    before = []
    after = []
    left = []
    first = []
    alignment = []
    for paragraph in paragraphs:
        value = paragraph_value(paragraph, "line_spacing")
        if hasattr(value, "pt"):
            value = ("pt", round(value.pt, 3))
        elif value is not None:
            value = ("multiple", round(float(value), 3))
        line_spacing.append(value)
        before.append(pt(paragraph_value(paragraph, "space_before")))
        after.append(pt(paragraph_value(paragraph, "space_after")))
        left.append(inches(paragraph_value(paragraph, "left_indent")))
        first.append(inches(paragraph_value(paragraph, "first_line_indent")))
        value = paragraph_value(paragraph, "alignment")
        alignment.append(None if value is None else int(value))
    return {
        "font": weighted_run_mode(paragraphs, "name"),
        "size_pt": weighted_run_mode(paragraphs, "size"),
        "line_spacing": mode(line_spacing),
        "space_before_pt": mode(before),
        "space_after_pt": mode(after),
        "left_indent_in": mode(left),
        "first_line_indent_in": mode(first),
        "alignment": mode(alignment),
    }


def document_groups(doc: Document) -> dict[str, Any]:
    paragraphs = doc.paragraphs
    first_drawing = next((index for index, p in enumerate(paragraphs) if has_drawing(p)), len(paragraphs))
    first_legend = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if re.match(r"^(?:Fig\.|Figure)\s+\d+\.", paragraph.text.strip())
        ),
        len(paragraphs),
    )
    legend_start = min(first_drawing, first_legend)
    reference_start = next(
        (index for index, p in enumerate(paragraphs) if p.style.name == "EndNote Bibliography"),
        legend_start,
    )
    title = next((p for p in paragraphs if p.text.strip()), None)
    abstract_index = next(
        (index for index, p in enumerate(paragraphs) if p.text.strip().upper() == "ABSTRACT"),
        min(reference_start, len(paragraphs)),
    )
    front = [p for p in paragraphs[1:abstract_index] if p.text.strip()]
    body = [
        p
        for p in paragraphs[abstract_index:reference_start]
        if p.text.strip() and not is_heading(p) and p.style.name != "EndNote Bibliography"
    ]
    headings = [p for p in paragraphs[:reference_start] if is_heading(p)]
    bibliography = [p for p in paragraphs if p.style.name == "EndNote Bibliography" and p.text.strip()]
    legends = [
        p
        for p in paragraphs[legend_start:]
        if p.text.strip() and not has_drawing(p) and not p.style.name.startswith("Heading")
    ]
    return {
        "title": paragraph_group_signature([title] if title is not None else []),
        "front_matter": paragraph_group_signature(front),
        "body": paragraph_group_signature(body),
        "headings": paragraph_group_signature(headings),
        "bibliography": paragraph_group_signature(bibliography),
        "figure_legends": paragraph_group_signature(legends),
    }


def package_signature(path: Path, doc: Document) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        footer_parts = {
            name: archive.read(name).decode("utf-8", errors="ignore")
            for name in names
            if re.fullmatch(r"word/footer\d+\.xml", name)
        }
        page_number_fields_by_footer = {
            name: len(
                re.findall(
                    r"<w:instrText[^>]*>[^<]*\bPAGE\b[^<]*</w:instrText>",
                    xml,
                    flags=re.I,
                )
            )
            for name, xml in footer_parts.items()
        }
        page_number_field_count = sum(page_number_fields_by_footer.values())
        media = {
            Path(name).name
            for name in names
            if name.startswith("word/media/") and not name.endswith("/")
        }
        referenced_media: set[str] = set()
        for name in names:
            if not name.endswith(".rels"):
                continue
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue
            for relationship in root:
                if relationship.attrib.get("Type", "").endswith("/image"):
                    referenced_media.add(Path(relationship.attrib.get("Target", "")).name)
    anchors = sum(len(paragraph._p.xpath(".//wp:anchor")) for paragraph in doc.paragraphs)
    sizes = [
        [round(shape.width.inches, 3), round(shape.height.inches, 3)]
        for shape in doc.inline_shapes
    ]
    usable_widths = [
        section.page_width - section.left_margin - section.right_margin
        for section in doc.sections
    ]
    full_width = sum(
        any(abs(shape.width - usable_width) <= 0.05 * 914400 for usable_width in usable_widths)
        for shape in doc.inline_shapes
    )
    within_page_width = sum(
        any(shape.width <= usable_width + 0.02 * 914400 for usable_width in usable_widths)
        for shape in doc.inline_shapes
    )
    return {
        "inline_figures": len(doc.inline_shapes),
        "anchored_figures": anchors,
        "inline_sizes_in": sizes,
        "full_text_width_figures": full_width,
        "figures_within_page_width": within_page_width,
        "comments_present": "word/comments.xml" in names,
        "tracked_changes_present": bool(
            re.search(r"<w:(?:ins|del)(?:\s|>)", document_xml)
        ),
        "page_number_field": page_number_field_count > 0,
        "page_number_field_count": page_number_field_count,
        "page_number_fields_by_footer": page_number_fields_by_footer,
        "no_duplicate_page_number_fields": page_number_field_count == 1,
        "orphaned_media": sorted(media - referenced_media),
    }


def signature(path: Path) -> dict[str, Any]:
    doc = Document(path)
    return {
        "section": section_signature(doc),
        "section_page_sizes_in": [
            [round(section.page_width.inches, 3), round(section.page_height.inches, 3)]
            for section in doc.sections
        ],
        "groups": document_groups(doc),
        "package": package_signature(path, doc),
    }


def close(left, right, tolerance: float = 0.02) -> bool:
    if left is None or right is None:
        return left is right
    if isinstance(left, (int, float)) and isinstance(right, (int, float)):
        return abs(float(left) - float(right)) <= tolerance
    return left == right


def compare_group(reference: dict[str, Any] | None, target: dict[str, Any] | None) -> dict[str, bool]:
    if reference is None:
        return {"reference_component_absent": True}
    if target is None:
        return {"target_component_present": False}
    keys = (
            "font",
            "size_pt",
            "line_spacing",
            "space_before_pt",
            "space_after_pt",
            "left_indent_in",
            "first_line_indent_in",
            "alignment",
        )
    comparisons = {}
    for key in keys:
        left = reference.get(key)
        right = target.get(key)
        if key in {
            "space_before_pt",
            "space_after_pt",
            "left_indent_in",
            "first_line_indent_in",
        }:
            left = 0.0 if left is None else left
            right = 0.0 if right is None else right
        comparisons[key] = close(left, right)
    return comparisons


def section_pages_match_reference(reference: dict[str, Any], target: dict[str, Any]) -> bool:
    """Allow the reference page size and its exact landscape rotation only."""
    reference_width = float(reference["section"]["page_width_in"])
    reference_height = float(reference["section"]["page_height_in"])
    allowed = (
        (reference_width, reference_height),
        (reference_height, reference_width),
    )
    return all(
        any(
            close(width, allowed_width) and close(height, allowed_height)
            for allowed_width, allowed_height in allowed
        )
        for width, height in target["section_page_sizes_in"]
    )


def find_paragraph_index(paragraphs, label: str) -> int | None:
    return next(
        (index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == label),
        None,
    )


def find_paragraph_prefix_index(paragraphs, prefix: str) -> int | None:
    return next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip().startswith(prefix)
        ),
        None,
    )


def has_section_break_before(paragraphs, index: int) -> bool:
    """Accept a next-page section break as an explicit page start."""
    return index > 0 and bool(paragraphs[index - 1]._p.xpath("./w:pPr/w:sectPr"))


def requested_structure_checks(args: argparse.Namespace, doc: Document) -> dict[str, Any]:
    paragraphs = doc.paragraphs
    checks: dict[str, Any] = {}
    if args.require_section_order:
        labels = [label.strip() for label in args.require_section_order.split(">") if label.strip()]
        positions = [find_paragraph_index(paragraphs, label) for label in labels]
        checks["required_section_order"] = (
            bool(labels)
            and all(position is not None for position in positions)
            and positions == sorted(positions)
            and len(set(positions)) == len(positions)
        )
        checks["required_section_positions"] = dict(zip(labels, positions))
    if args.require_paragraph_prefix_order:
        prefixes = [
            prefix.strip()
            for prefix in args.require_paragraph_prefix_order.split(">")
            if prefix.strip()
        ]
        positions = [
            find_paragraph_prefix_index(paragraphs, prefix)
            for prefix in prefixes
        ]
        checks["required_paragraph_prefix_order"] = (
            bool(prefixes)
            and all(position is not None for position in positions)
            and positions == sorted(positions)
            and len(set(positions)) == len(positions)
        )
        checks["required_paragraph_prefix_positions"] = dict(
            zip(prefixes, positions)
        )
    if args.require_page_break_before:
        checks["required_page_breaks"] = {
            label: (
                (index := find_paragraph_index(paragraphs, label)) is not None
                and (
                    bool(paragraph_value(paragraphs[index], "page_break_before"))
                    or has_section_break_before(paragraphs, index)
                )
            )
            for label in args.require_page_break_before
        }
    drawing_indices = [
        index for index, paragraph in enumerate(paragraphs) if has_drawing(paragraph)
    ]
    if args.figures_after_heading:
        heading_index = find_paragraph_index(paragraphs, args.figures_after_heading)
        checks["figures_after_requested_heading"] = (
            heading_index is not None
            and all(index > heading_index for index in drawing_indices)
        )
        checks["requested_figure_heading_position"] = heading_index
    if args.require_figure_legend_pairs:
        pair_count = 0
        legend_pattern = re.compile(r"^(?:Supplementary\s+)?(?:Fig\.|Figure)\s+S?\d+\.")
        for position, drawing_index in enumerate(drawing_indices):
            next_drawing = (
                drawing_indices[position + 1]
                if position + 1 < len(drawing_indices)
                else len(paragraphs)
            )
            if any(
                legend_pattern.match(paragraph.text.strip())
                for paragraph in paragraphs[drawing_index + 1:next_drawing]
            ):
                pair_count += 1
        checks["figure_legend_pair_count"] = pair_count
        checks["every_figure_has_adjacent_legend"] = (
            pair_count == len(drawing_indices)
            and (
                args.expected_figures is None
                or pair_count == args.expected_figures
            )
        )
    return checks


def main() -> int:
    args = parse_args()
    reference = signature(args.reference)
    target = signature(args.target)
    target_doc = Document(args.target)

    comparisons: dict[str, Any] = {
        "section": {
            key: close(reference["section"][key], target["section"][key])
            for key in reference["section"]
        },
        "section_page_geometry": {
            "all_sections_match_reference_size_or_rotation":
                section_pages_match_reference(reference, target),
        },
        "groups": {
            key: compare_group(reference["groups"][key], target["groups"][key])
            for key in reference["groups"]
        },
        "figures": {
            "expected_inline_count": (
                True
                if args.expected_figures is None
                else target["package"]["inline_figures"] == args.expected_figures
            ),
            "no_anchored_figures": target["package"]["anchored_figures"] == 0,
            "all_inline_figures_fit_page_width": (
                target["package"]["inline_figures"] == 0
                or target["package"]["figures_within_page_width"]
                == target["package"]["inline_figures"]
            ),
            "comments_absent": not target["package"]["comments_present"],
            "tracked_changes_absent": not target["package"]["tracked_changes_present"],
            "page_number_field_present": target["package"]["page_number_field"],
            "single_page_number_field": target["package"]["no_duplicate_page_number_fields"],
            "no_orphaned_media": not target["package"]["orphaned_media"],
        },
        "requested_structure": requested_structure_checks(args, target_doc),
    }
    flat = []
    for component in comparisons.values():
        if isinstance(component, dict):
            for value in component.values():
                if isinstance(value, dict):
                    flat.extend(item for item in value.values() if isinstance(item, bool))
                else:
                    if isinstance(value, bool):
                        flat.append(value)
    status = "pass" if all(flat) else "fail"
    payload = {
        "status": status,
        "reference": str(args.reference),
        "target": str(args.target),
        "reference_signature": reference,
        "target_signature": target,
        "comparisons": comparisons,
    }
    rendered = json.dumps(payload, indent=2, ensure_ascii=False) + "\n"
    if args.json:
        args.json.parent.mkdir(parents=True, exist_ok=True)
        args.json.write_text(rendered, encoding="utf-8")
    print(rendered, end="")
    return 0 if status == "pass" else 1


if __name__ == "__main__":
    sys.exit(main())
